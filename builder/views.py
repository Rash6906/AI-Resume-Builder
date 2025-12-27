import re
import json
import pdfplumber
import docx
import io
import os
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.http import HttpResponse, JsonResponse
from django.template.loader import get_template
from django.contrib.auth.forms import UserCreationForm
from django.contrib.auth import login
from django.forms import inlineformset_factory
from django.views.decorators.csrf import csrf_exempt
from django.conf import settings
from xhtml2pdf import pisa
from huggingface_hub import InferenceClient

from .models import Resume, Experience, Hobby, Education, Skill, Project, Certification
from .forms import ResumeForm, EducationFormSet, ExperienceFormSet, SkillFormSet, ProjectFormSet, CertificationFormSet, HobbyFormSet
from django.shortcuts import render
from django.contrib.auth.decorators import login_required
from .models import Resume
import random
import io
from docx import Document
from pdfminer.high_level import extract_text as extract_pdf_text
from django.core.files.storage import FileSystemStorage
import io
from docx import Document
from pdfminer.high_level import extract_text as extract_pdf_text
from django.shortcuts import render, redirect
from django.contrib import messages
# --- AI Configuration ---
import os
from dotenv import load_dotenv
load_dotenv()

HF_TOKEN = os.getenv("HF_TOKEN")
MODEL_ID = "meta-llama/Meta-Llama-3-8B-Instruct"

# Use the token from the environment variable
client = InferenceClient(token=HF_TOKEN)

# --- Helper: Formset Factory ---
def get_resume_formsets(extra_val=1):
    return {
        'EduFS': inlineformset_factory(Resume, Education, fields='__all__', extra=extra_val, can_delete=True),
        'ExpFS': inlineformset_factory(Resume, Experience, fields='__all__', extra=extra_val, can_delete=True),
        'SkillFS': inlineformset_factory(Resume, Skill, fields='__all__', extra=extra_val, can_delete=True),
        'ProjFS': inlineformset_factory(Resume, Project, fields='__all__', extra=extra_val, can_delete=True),
        'CertFS': inlineformset_factory(Resume, Certification, fields='__all__', extra=extra_val, can_delete=True),
        'HobbyFS': inlineformset_factory(Resume, Hobby, fields='__all__', extra=extra_val, can_delete=True),
    }

# --- AI Suggestion Logic ---
@csrf_exempt
def ai_suggest(request):
    if request.method == "POST":
        try:
            data = json.loads(request.body)
            raw_prompt = data.get("prompt", "")
            response = client.chat_completion(
                model=MODEL_ID,
                messages=[
                    {"role": "system", "content": "You are a professional resume writer."},
                    {"role": "user", "content": f"Write a professional summary for: {raw_prompt}"}
                ],
                max_tokens=400,
                temperature=0.8
            )
            suggestion = response.choices[0].message.content.strip()
            return JsonResponse({"suggestion": suggestion.replace('"', '').replace('**', '')})
        except Exception as e:
            return JsonResponse({"suggestion": "Error generating content."}, status=500)
    return JsonResponse({"error": "Invalid request"}, status=400)

# --- Main CRUD Views ---

def home(request):
    if request.user.is_authenticated:
        resumes = Resume.objects.filter(user=request.user)
        return render(request, 'builder/home.html', {'resumes': resumes})
    return render(request, 'builder/landing.html')

@login_required
def create_resume(request):
    factories = get_resume_formsets(extra_val=1)
    parsed_data = request.session.pop('parsed_resume_data', None)

    # Standardized 3-letter prefixes
    prefix_map = {
        'EduFS': 'edu', 'ExpFS': 'exp', 'SkillFS': 'ski',
        'ProjFS': 'pro', 'CertFS': 'cer', 'HobbyFS': 'hob'
    }

    if request.method == "POST":
        form = ResumeForm(request.POST, request.FILES)
        formsets = [factories[name](request.POST, prefix=p) for name, p in prefix_map.items()]
        
        if form.is_valid() and all(fs.is_valid() for fs in formsets):
            resume = form.save(commit=False)
            resume.user = request.user
            resume.save()
            for fs in formsets:
                fs.instance = resume
                fs.save()
            return redirect('preview_resume', pk=resume.pk)
    else:
        form = ResumeForm(initial=parsed_data['main'] if parsed_data else {})
        context = {
            'form': form,
            'edu_fs': factories['EduFS'](prefix='edu'),
            'exp_fs': factories['ExpFS'](prefix='exp'),
            'ski_fs': factories['SkillFS'](prefix='ski'),
            'pro_fs': factories['ProjFS'](prefix='pro'),
            'cer_fs': factories['CertFS'](prefix='cer'),
            'hob_fs': factories['HobbyFS'](prefix='hob'),
        }
    return render(request, 'builder/resume_form.html', context)

@login_required
def edit_resume(request, pk):
    resume = get_object_or_404(Resume, pk=pk, user=request.user)
    factories = get_resume_formsets(extra_val=0)
    
    if request.method == "POST":
        form = ResumeForm(request.POST, request.FILES, instance=resume)
        formsets = {
            'edu_fs': factories['EduFS'](request.POST, instance=resume, prefix='edu'),
            'exp_fs': factories['ExpFS'](request.POST, instance=resume, prefix='exp'),
            'skill_fs': factories['SkillFS'](request.POST, instance=resume, prefix='ski'),
            'proj_fs': factories['ProjFS'](request.POST, instance=resume, prefix='pro'),
            'cert_fs': factories['CertFS'](request.POST, instance=resume, prefix='cer'),
            'hobby_fs': factories['HobbyFS'](request.POST, instance=resume, prefix='hob'),
        }

        if form.is_valid() and all(fs.is_valid() for fs in formsets.values()):
            form.save()
            for fs in formsets.values():
                fs.save()
            return redirect('home')
    else:
        form = ResumeForm(instance=resume)
        formsets = {
            'edu_fs': factories['EduFS'](instance=resume, prefix='edu'),
            'exp_fs': factories['ExpFS'](instance=resume, prefix='exp'),
            'skill_fs': factories['SkillFS'](instance=resume, prefix='ski'),
            'proj_fs': factories['ProjFS'](instance=resume, prefix='pro'),
            'cert_fs': factories['CertFS'](instance=resume, prefix='cer'),
            'hobby_fs': factories['HobbyFS'](instance=resume, prefix='hob'),
        }

    context = {'form': form, 'edit_mode': True, **formsets}
    return render(request, 'builder/resume_form.html', context)

# --- Parsing and Upload ---

@login_required
def upload_existing_resume(request):
    if request.method == "POST" and request.FILES.get('resume_file'):
        file = request.FILES['resume_file']
        ext = file.name.split('.')[-1].lower()
        text = ""

        if ext == 'pdf':
            with pdfplumber.open(file) as pdf:
                text = "\n".join([p.extract_text() for p in pdf.pages if p.extract_text()])
        elif ext in ['doc', 'docx']:
            doc = docx.Document(file)
            text = "\n".join([p.text for p in doc.paragraphs])

        # AI Structured Extraction
        ai_prompt = f"""
        Extract data from the following resume text into JSON. 
        Structure: 
        {{
            "main": {{"full_name": "", "email": "", "phone": "", "summary": ""}},
            "edu": [{{ "institution": "", "degree": "", "year": "" }}],
            "exp": [{{ "company": "", "position": "", "description": "" }}],
            "ski": [{{ "name": "" }}],
            "pro": [{{ "title": "", "description": "" }}],
            "cer": [{{ "name": "" }}],
            "hob": [{{ "name": "" }}]
        }}
        Text: {text[:2000]}
        """
        try:
            response = client.chat_completion(
                model=MODEL_ID,
                messages=[{"role": "user", "content": ai_prompt}],
                max_tokens=1000
            )
            # Find JSON block in AI response
            raw_content = response.choices[0].message.content
            json_match = re.search(r'\{.*\}', raw_content, re.DOTALL)
            if json_match:
                parsed_data = json.loads(json_match.group())
                request.session['parsed_resume_data'] = parsed_data
        except Exception as e:
            print(f"AI Parse Error: {e}")

        return redirect('create_resume')
    return redirect('home')

# --- Preview & PDF Generation ---
# In views.py
@login_required
def preview_resume(request, pk):
    resume = get_object_or_404(Resume, pk=pk)
    style = request.GET.get('style', 'ats')
    
    templates = {
        'ats': 'builder/resume_ats.html',
        'modern': 'builder/resume_modern.html',
        'classic': 'builder/resume_classic.html',
        'creative_elite': 'builder/creative_elite.html', # This key MUST match the ?style= link
    }
    
    template_name = templates.get(style, 'builder/resume_ats.html')
    return render(request, template_name, {'resume': resume, 'current_style': style})
@login_required
def download_resume_pdf(request, pk):
    resume = get_object_or_404(Resume, pk=pk)
    style = request.GET.get('style', 'classic')
    theme_color = request.GET.get('color', '#6610f2')
    print(f"DEBUG: User requested style: {style}")
    # Map styles to specific PDF templates
    template_map = {
        'creative': 'builder/pdf_creative.html',
        'modern': 'builder/pdf_modern.html',
        'ats': 'builder/pdf_ats.html',
        'classic': 'builder/pdf_template.html'
    }
    
    template_path = template_map.get(style, 'builder/pdf_template.html')
    context = {'resume': resume, 'current_style': style}
    
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{resume.full_name}_resume.pdf"'
    
    context = {
        'resume': resume,
        'theme_color': theme_color,
    }

    template = get_template(template_path)
    html = template.render(context)
    
    pisa_status = pisa.CreatePDF(html, dest=response, link_callback=link_callback)
    if pisa_status.err:
        return HttpResponse('We had some errors <pre>' + html + '</pre>')
    return response
def link_callback(uri, rel):
    """
    Convert HTML URIs to absolute system paths so xhtml2pdf can access those files
    """
    import os
    from django.conf import settings

    # Use settings to get the root paths
    s_url = settings.STATIC_URL     # e.g. '/static/'
    s_root = settings.STATIC_ROOT   # physical path to static folder
    m_url = settings.MEDIA_URL      # e.g. '/media/'
    m_root = settings.MEDIA_ROOT    # physical path to media folder

    if uri.startswith(m_url):
        path = os.path.join(m_root, uri.replace(m_url, ""))
    elif uri.startswith(s_url):
        path = os.path.join(s_root, uri.replace(s_url, ""))
    else:
        # If it's already an absolute path, just return it
        return uri

    # make sure that file exists
    if not os.path.isfile(path):
        # If the file isn't found, xhtml2pdf will attempt to download it 
        # but this often fails, so we return the original uri as a fallback
        return uri
        
    return path

# --- Authentication ---
@login_required
def register(request):
    if request.method == 'POST':
        form = UserCreationForm(request.POST)
        if form.is_valid():
            login(request, form.save())
            return redirect('home')
    return render(request, 'registration/register.html', {'form': UserCreationForm()})

@login_required
def delete_resume(request, pk):
    get_object_or_404(Resume, pk=pk, user=request.user).delete()
    return redirect('home')

# builder/views

@login_required
def job_matcher(request):
    resumes = Resume.objects.filter(user=request.user)
    
    # Initialize variables for the template
    score, feedback = None, ""
    matched_keywords, missing_keywords = [], []
    suggested_summary, action_items = "", []
    job_desc_input, selected_resume_id = "", None

    # Words to ignore to prevent score inflation (e.g., 'web' or 'developer' won't give 100%)
    STOP_WORDS = {
        'developer', 'engineer', 'manager', 'junior', 'senior', 'work', 
        'experience', 'skills', 'requirements', 'responsibilities', 'description',
        'with', 'from', 'that', 'this', 'team', 'company', 'role'
    }

    if request.method == "POST":
        job_desc_input = request.POST.get('job_description', '')
        selected_resume_id = request.POST.get('resume_id')
        
        selected_resume = get_object_or_404(Resume, id=selected_resume_id, user=request.user)
        
        # 1. Grab all text fields dynamically from your Resume model
        all_text_parts = []
        for field in selected_resume._meta.fields:
            value = getattr(selected_resume, field.name)
            if isinstance(value, str):
                all_text_parts.append(str(value).lower())
        resume_text = " ".join(all_text_parts)
        
        # 2. Extract and Filter Words
        job_words = set(re.findall(r'\w+', job_desc_input.lower()))
        resume_words = set(re.findall(r'\w+', resume_text))
        
        # Focus only on technical keywords (Length > 3 and not a Stop Word)
        meaningful_job_words = {word for word in job_words if len(word) > 3 and word not in STOP_WORDS}
        
        if meaningful_job_words:
            matches = meaningful_job_words.intersection(resume_words)
            missing = meaningful_job_words.difference(resume_words)
            
            matched_keywords = sorted(list(matches))
            missing_keywords = sorted(list(missing))
            
            # 3. Realistic Scoring (Matches / Total Meaningful)
            score = int((len(matches) / len(meaningful_job_words)) * 100)
            
            # 4. Generate AI-Style Advice
            if matched_keywords:
                suggested_summary = f"Experienced professional with a background in {', '.join(matched_keywords[:3])}. " \
                                   f"Eager to apply expertise in {matched_keywords[-1] if len(matched_keywords)>3 else 'software development'} " \
                                   f"to drive results in this role."
            
            action_items = [
                f"Incorporate '{missing_keywords[0]}' and '{missing_keywords[1]}' into your Skills section." if len(missing_keywords) > 1 else "Skills look good!",
                "Quantify your impact with numbers (e.g., 'Reduced costs by 15%').",
                "Ensure your summary specifically mentions the job title."
            ]
        else:
            score = 0

        # Feedback Logic
        if score > 80: feedback = "Excellent Match! Your resume is highly compatible."
        elif score > 50: feedback = "Good Match. A few adjustments could make it perfect."
        else: feedback = "Low Match. Follow the improvement plan below."

    return render(request, 'builder/job_matcher.html', {
        'resumes': resumes,
        'score': score,
        'feedback': feedback,
        'matched_keywords': matched_keywords,
        'missing_keywords': missing_keywords[:12],  # Show top 12 missing
        'job_description': job_desc_input,
        'selected_resume_id': selected_resume_id,
        'suggested_summary': suggested_summary,
        'action_items': action_items
    })

def ats_format_checker(resume_text):
    issues = []
    # Check for standard sections
    sections = ['experience', 'education', 'skills', 'contact']
    for s in sections:
        if s not in resume_text.lower():
            issues.append(f"Missing standard section header: {s.capitalize()}")
            
    # Check for tables/images (if processing a file)
    # If the text density is too low compared to file size, it might have images
    return issues

def perform_ats_scan(text):
    """
    Analyzes resume text for structural 'ATS-friendliness'.
    Returns a score and a list of specific formatting issues.
    """
    text_lower = text.lower()
    issues = []
    format_score = 100

    # 1. Check for Standard Section Headers
    # If headers are missing, ATS might not know where 'Skills' end and 'Work' begins.
    required_sections = {
        'experience': ['experience', 'work history', 'employment'],
        'education': ['education', 'academic'],
        'skills': ['skills', 'core competencies', 'technologies'],
        'contact': ['email', 'phone', 'contact', 'address']
    }

    for section, keywords in required_sections.items():
        if not any(k in text_lower for k in keywords):
            issues.append(f"Missing or non-standard '{section.capitalize()}' header.")
            format_score -= 15

    # 2. Check for Contact Info (Email/Phone)
    # ATS needs these to create your profile.
    email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    phone_pattern = r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
    
    if not re.search(email_pattern, text):
        issues.append("Email address not detected or unreadable.")
        format_score -= 20
    if not re.search(phone_pattern, text):
        issues.append("Phone number not detected (use standard format like 123-456-7890).")
        format_score -= 10

    # 3. Check for "Machine Readability"
    # If the text is too short but the resume is long, it might be an 'image-only' PDF.
    if len(text.split()) < 50:
        issues.append("Low text count: Ensure your resume is not an image/scan.")
        format_score -= 30

    return max(format_score, 0), issues

def ats_compliance_checker(request, resume_id):
    resume = get_object_or_404(Resume, id=resume_id, user=request.user)
    
    # 1. Gather text
    all_text = []
    for field in resume._meta.fields:
        val = getattr(resume, field.name)
        if isinstance(val, str): all_text.append(str(val))
    resume_text = " ".join(all_text)
    
    # 2. Run the Compliance Scan
    # We look for common "Robot" deal-breakers
    issues = []
    score = 100

    # Check for Contact Info
    if not re.search(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', resume_text):
        issues.append("Email not found. Robots won't know how to reach you.")
        score -= 20

    # Check for Standard Headers
    headers = ['experience', 'education', 'skills', 'projects']
    for h in headers:
        if h not in resume_text.lower():
            issues.append(f"Standard '{h.capitalize()}' header missing.")
            score -= 15

    # Check for Quantifiable Metrics (The % and $ signs recruiters love)
    if not re.search(r'[%$]|\d+', resume_text):
        issues.append("Lack of metrics. Try adding numbers (e.g., 'Increased sales by 20%').")
        score -= 10

    return render(request, 'builder/ats_report.html', {
        'resume': resume,
        'ats_score': max(score, 0),
        'issues': issues
    })

@login_required
def ats_scanner(request, resume_id):
    resume = get_object_or_404(Resume, id=resume_id)
    
    # 1. Collect all related data into strings
    # We use .all() on the related_names you defined in models.py
    skills = ", ".join([s.skill_name for s in resume.skills.all()])
    
    exp_list = [f"{e.role} at {e.company}: {e.description}" for e in resume.experiences.all()]
    experience = " ".join(exp_list)
    
    edu_list = [f"{ed.degree} from {ed.institution}" for ed in resume.education_details.all()]
    education = " ".join(edu_list)

    # 2. Combine into one big block of text for the scanner
    resume_text = f"{resume.full_name} {resume.objective} {experience} {skills} {education}"

    # 3. Now run your scoring logic on resume_text
    # ... (rest of your scanner logic) ...
    
    # Example: Simple score based on existence of text
    score = 0
    if len(skills) > 5: score += 30
    if len(experience) > 20: score += 40
    if len(education) > 10: score += 30

    return render(request, 'builder/ats_results.html', {
        'resume': resume,
        'score': score,
        'raw_text': resume_text
    })
@login_required
def ats_scanner_select(request):
    """Bridge view to handle the dropdown selection from the home page."""
    resume_id = request.GET.get('resume_id')
    if resume_id:
        # Redirects to the actual scanner view we built earlier
        return redirect('ats_scanner', resume_id=resume_id)
    return redirect('home')


def ats_scanner_upload(request):
    if request.method == 'POST' and request.FILES.get('resume_file'):
        uploaded_file = request.FILES['resume_file']
        file_name = uploaded_file.name
        resume_text = ""

        try:
            # 1. Extraction
            if file_name.endswith('.pdf'):
                resume_text = extract_pdf_text(io.BytesIO(uploaded_file.read()))
            elif file_name.endswith('.docx'):
                doc = Document(uploaded_file)
                resume_text = "\n".join([para.text for para in doc.paragraphs])
            
            if not resume_text.strip():
                raise ValueError("Could not extract text from file.")

            # 2. Refined Audit Scoring
            score = 0
            results = []

            # A. Contact Info Check (Email & Phone)
            has_email = bool(re.search(r'[\w\.-]+@[\w\.-]+', resume_text))
            has_phone = bool(re.search(r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}', resume_text))
            if has_email and has_phone:
                score += 25
                results.append({'label': 'Contact Info', 'status': True, 'desc': 'Found both email and phone number.'})
            else:
                results.append({'label': 'Contact Info', 'status': False, 'desc': 'Missing or poorly formatted email/phone.'})

            # B. Experience Date Check (Looks for "2018 - 2022" or "Present")
            has_dates = bool(re.search(r'(\d{4}|Present)', resume_text))
            if 'experience' in resume_text.lower() and has_dates:
                score += 25
                results.append({'label': 'Work History', 'status': True, 'desc': 'Experience found with valid timeline dates.'})
            else:
                results.append({'label': 'Work History', 'status': False, 'desc': 'Timeline dates or Experience header missing.'})

            # C. Education Audit
            if any(word in resume_text.lower() for word in ['education', 'university', 'college', 'degree']):
                score += 25
                results.append({'label': 'Education', 'status': True, 'desc': 'Academic background detected.'})
            else:
                results.append({'label': 'Education', 'status': False, 'desc': 'No Education section identified.'})

            # D. Skills & Formatting (Bullet points detection)
            has_bullets = resume_text.count('•') > 3 or resume_text.count('- ') > 3
            if 'skills' in resume_text.lower() and has_bullets:
                score += 25
                results.append({'label': 'Skills & Format', 'status': True, 'desc': 'Skills found with clean bullet-point formatting.'})
            else:
                results.append({'label': 'Skills & Format', 'status': False, 'desc': 'Skills section missing or lacks bullet points.'})

            return render(request, 'builder/ats_results.html', {
                'score': score,
                'results': results,
                'raw_text': resume_text[:2000],
                'file_name': file_name
            })

        except Exception as e:
            messages.error(request, f"Scan failed: {str(e)}")
            return redirect('home')

    return redirect('home')

def calculate_ats_score(text):
    score = 0
    issues = []
    
    # Example: Check for standard headers
    headers = ['experience', 'education', 'skills', 'projects']
    for header in headers:
        if header in text.lower():
            score += 15
        else:
            issues.append(f"Missing section: {header.capitalize()}")
            
    # Example: Check for contact info
    if "@" in text:
        score += 20
    else:
        issues.append("No email address detected.")

    return score, issues